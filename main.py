import os
import json
import sqlite3
import datetime
import gradio as gr
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt
import tempfile
import shutil

# Constants
OLLAMA_API_URL = "http://localhost:11434/api/generate"
DB_PATH = "job_applications.db"
RESUME_FOLDER = "resumes"
COVER_LETTER_FOLDER = "cover_letters"

# Ensure folders exist
os.makedirs(RESUME_FOLDER, exist_ok=True)
os.makedirs(COVER_LETTER_FOLDER, exist_ok=True)

# Database setup
def setup_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Create job applications table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS job_applications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company_name TEXT NOT NULL,
        position TEXT NOT NULL,
        date_applied TEXT NOT NULL,
        job_description TEXT,
        status TEXT NOT NULL,
        salary_info TEXT,
        contact_info TEXT,
        application_url TEXT,
        notes TEXT,
        resume_path TEXT,
        cover_letter_path TEXT,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )
    ''')
    
    # Create user information table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS user_information (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        full_name TEXT NOT NULL,
        address TEXT,
        phone TEXT,
        email TEXT NOT NULL,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )
    ''')
    
    conn.commit()
    conn.close()

# Ollama integration
def check_ollama_availability():
    try:
        response = requests.get("http://localhost:11434/api/version")
        response.raise_for_status()
        return True, None
    except Exception as e:
        return False, str(e)

def generate_text_with_ollama(prompt, model="qwen2.5:14b"):
    # Check if Ollama is available
    available, error = check_ollama_availability()
    if not available:
        return f"Error: Ollama is not available. Please make sure Ollama is running. Details: {error}"
    
    try:
        response = requests.post(
            OLLAMA_API_URL,
            json={
                "model": model,
                "prompt": prompt,
                "stream": False
            }
        )
        response.raise_for_status()
        return response.json().get("response", "")
    except Exception as e:
        error_msg = str(e)
        if "404" in error_msg:
            return "Error: The API endpoint was not found. Please check if Ollama is running correctly and the API URL is correct."
        elif "Connection refused" in error_msg:
            return "Error: Connection refused. Please make sure Ollama is running."
        else:
            return f"Error generating text with Ollama: {error_msg}"

# Document processing functions
def create_custom_resume(base_resume_path, job_description, output_path=None):
    # Generate tailored content using Ollama
    prompt = f"""
    I have a job description and need to customize my resume for it.
    
    Job Description:
    {job_description}
    
    Please analyze this job description and provide specific suggestions on how I should tailor my resume.
    Focus on:
    1. Skills to emphasize
    2. Experience to highlight
    3. Achievements that would be most relevant
    4. Keywords to include
    
    Format your response as specific, actionable bullet points I can use to modify my resume.
    """
    
    tailoring_suggestions = generate_text_with_ollama(prompt)
    
    # Load the base resume
    doc = Document(base_resume_path)
    
    # Add tailoring suggestions at the end
    # Check if the style exists in the document
    try:
        doc.add_heading("Tailoring Suggestions", level=1)
    except KeyError:
        # If the style doesn't exist, use a paragraph with manual formatting
        p = doc.add_paragraph("Tailoring Suggestions")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(16)  # Approximate size for Heading 1
    
    doc.add_paragraph(tailoring_suggestions)
    
    # Save the customized resume
    if not output_path:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(RESUME_FOLDER, f"custom_resume_{timestamp}.docx")
    
    doc.save(output_path)
    return output_path, tailoring_suggestions

def generate_cover_letter(job_description, company_name, position, resume_path, output_path=None):
    # Extract resume content
    doc = Document(resume_path)
    resume_text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    
    # Get user information
    user_info = get_user_information()
    
    # Generate cover letter using Ollama
    prompt = f"""
    Write a professional cover letter for a {position} position at {company_name}.
    
    Job Description:
    {job_description}
    
    My Resume:
    {resume_text}
    
    The cover letter should:
    1. Be professionally formatted
    2. Highlight relevant skills and experience from my resume that match the job requirements
    3. Show enthusiasm for the role and company
    4. Include a strong opening and closing
    5. Be approximately 300-400 words
    6. Only mention skills and experience that are actually in my resume
    7. Specifically mention the company name ({company_name}) and position ({position})
    8. Reference specific requirements or qualifications from the job description
    
    Write the complete cover letter text, ready to be used.
    """
    
    cover_letter_text = generate_text_with_ollama(prompt)
    
    # Create a new document
    doc = Document()
    
    # Add user information at the top if available
    if user_info["full_name"]:
        doc.add_paragraph(user_info["full_name"])
        if user_info["address"]:
            doc.add_paragraph(user_info["address"])
        if user_info["phone"]:
            doc.add_paragraph(user_info["phone"])
        if user_info["email"]:
            doc.add_paragraph(user_info["email"])
    else:
        doc.add_paragraph("Your Name")
        doc.add_paragraph("Your Address")
        doc.add_paragraph("Your Phone")
        doc.add_paragraph("Your Email")
    
    # Add date
    doc.add_paragraph(datetime.datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()
    
    # Add company info
    doc.add_paragraph("Hiring Manager")
    doc.add_paragraph(f"{company_name}")
    doc.add_paragraph("Company Address")
    doc.add_paragraph("City, State ZIP")
    doc.add_paragraph()
    
    # Add greeting
    doc.add_paragraph("Dear Hiring Manager,")
    
    # Add cover letter content
    paragraphs = cover_letter_text.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    # Add closing
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(user_info["full_name"] if user_info["full_name"] else "Your Name")
    
    # Save the cover letter
    if not output_path:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(COVER_LETTER_FOLDER, f"cover_letter_{timestamp}.docx")
    
    doc.save(output_path)
    return output_path, cover_letter_text

# User information operations
def save_user_information(full_name, address, phone, email):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Check if user information already exists
    cursor.execute('SELECT COUNT(*) FROM user_information')
    count = cursor.fetchone()[0]
    
    if count > 0:
        # Update existing record
        cursor.execute('''
        UPDATE user_information
        SET full_name = ?, address = ?, phone = ?, email = ?, updated_at = ?
        WHERE id = 1
        ''', (full_name, address, phone, email, now))
    else:
        # Insert new record
        cursor.execute('''
        INSERT INTO user_information (
            full_name, address, phone, email, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?)
        ''', (full_name, address, phone, email, now, now))
    
    conn.commit()
    conn.close()
    return True

def get_user_information():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM user_information LIMIT 1')
    row = cursor.fetchone()
    
    conn.close()
    return dict(row) if row else {"full_name": "", "address": "", "phone": "", "email": ""}

# Database operations
def add_job_application(company_name, position, date_applied, job_description, status, 
                       salary_info=None, contact_info=None, application_url=None, notes=None,
                       resume_path=None, cover_letter_path=None):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    cursor.execute('''
    INSERT INTO job_applications (
        company_name, position, date_applied, job_description, status,
        salary_info, contact_info, application_url, notes,
        resume_path, cover_letter_path, created_at, updated_at
    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        company_name, position, date_applied, job_description, status,
        salary_info, contact_info, application_url, notes,
        resume_path, cover_letter_path, now, now
    ))
    
    job_id = cursor.lastrowid
    conn.commit()
    conn.close()
    
    return job_id

def update_job_application(job_id, **kwargs):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Add updated_at timestamp
    kwargs['updated_at'] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Build the SET clause for the SQL query
    set_clause = ", ".join([f"{key} = ?" for key in kwargs.keys()])
    values = list(kwargs.values())
    values.append(job_id)  # For the WHERE clause
    
    cursor.execute(f'''
    UPDATE job_applications
    SET {set_clause}
    WHERE id = ?
    ''', values)
    
    conn.commit()
    conn.close()

def get_all_job_applications():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM job_applications ORDER BY date_applied DESC')
    rows = cursor.fetchall()
    
    applications = []
    for row in rows:
        applications.append(dict(row))
    
    conn.close()
    return applications

def get_job_application(job_id):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM job_applications WHERE id = ?', (job_id,))
    row = cursor.fetchone()
    
    conn.close()
    return dict(row) if row else None

def delete_job_application(job_id):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Get file paths before deletion
    cursor.execute('SELECT resume_path, cover_letter_path FROM job_applications WHERE id = ?', (job_id,))
    row = cursor.fetchone()
    
    # Delete from database
    cursor.execute('DELETE FROM job_applications WHERE id = ?', (job_id,))
    
    conn.commit()
    conn.close()
    
    # Delete associated files if they exist
    if row:
        resume_path, cover_letter_path = row
        if resume_path and os.path.exists(resume_path):
            os.remove(resume_path)
        if cover_letter_path and os.path.exists(cover_letter_path):
            os.remove(cover_letter_path)
    
    return True

# Custom Theme
def create_custom_theme():
    return gr.themes.Soft(
        primary_hue="indigo",
        secondary_hue="blue",
        neutral_hue="slate",
        font=["Inter", "ui-sans-serif", "system-ui", "sans-serif"],
    )

# Gradio UI
def create_ui():
    # Set a custom theme
    theme = create_custom_theme()
    # CSS for custom styling
    custom_css = """
    .gradio-container {
        max-width: 1200px !important;
    }
    
    h1 {
        text-align: center;
        margin-bottom: 1.5rem;
        color: #4f46e5;
        font-weight: 700;
        font-size: 2.5rem;
    }
    
    .app-header {
        display: flex;
        align-items: center;
        justify-content: center;
        margin-bottom: 1rem;
        gap: 0.5rem;
    }
    
    .app-header svg {
        width: 32px;
        height: 32px;
    }
    
    .tab-header {
        font-size: 1.25rem;
        font-weight: 600;
        margin-bottom: 1rem;
        color: #4f46e5;
        border-bottom: 2px solid #e5e7eb;
        padding-bottom: 0.5rem;
    }
    
    .info-box {
        background-color: #f0f9ff;
        border-left: 4px solid #3b82f6;
        padding: 1rem;
        margin-bottom: 1rem;
        border-radius: 0.375rem;
    }
    
    .success-message {
        color: #047857;
        font-weight: 500;
    }
    
    .error-message {
        color: #b91c1c;
        font-weight: 500;
    }
    
    .footer {
        text-align: center;
        margin-top: 2rem;
        padding-top: 1rem;
        border-top: 1px solid #e5e7eb;
        color: #6b7280;
        font-size: 0.875rem;
    }
    """
    
    # Helper function to refresh applications list
    def refresh_applications_list():
        print("Refreshing applications list...")
        try:
            applications = get_all_job_applications()
            print(f"Retrieved {len(applications)} applications")
            if applications:
                df = pd.DataFrame(applications)
                # Format the dataframe for display
                display_df = df[['id', 'company_name', 'position', 'date_applied', 'status']].copy()
                return display_df
            return pd.DataFrame(columns=['id', 'company_name', 'position', 'date_applied', 'status'])
        except Exception as e:
            print(f"Error in refresh_applications_list: {str(e)}")
            return pd.DataFrame(columns=['id', 'company_name', 'position', 'date_applied', 'status'])
    
    # Helper function to save uploaded resume
    def save_uploaded_resume(file):
        print("In save_uploaded_resume function...")
        if file is None:
            print("No file provided")
            return None
        
        try:
            # Create a unique filename
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"base_resume_{timestamp}.docx"
            filepath = os.path.join(RESUME_FOLDER, filename)
            
            print(f"File info: {file}")
            print(f"Saving to: {filepath}")
            
            # Save the uploaded file
            shutil.copy(file.name, filepath)
            print(f"File saved successfully to {filepath}")
            return filepath
        except Exception as e:
            print(f"Error in save_uploaded_resume: {str(e)}")
            return None
    
    # Add Job Application Tab
    with gr.Blocks(theme=theme) as app:
        with gr.Row(elem_classes="app-header"):
            gr.Markdown("""
            <h1>
                <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="currentColor">
                    <path d="M4.5 6.375a4.125 4.125 0 118.25 0 4.125 4.125 0 01-8.25 0zM14.25 8.625a3.375 3.375 0 116.75 0 3.375 3.375 0 01-6.75 0zM1.5 19.125a7.125 7.125 0 0114.25 0v.003l-.001.119a.75.75 0 01-.363.63 13.067 13.067 0 01-6.761 1.873c-2.472 0-4.786-.684-6.76-1.873a.75.75 0 01-.364-.63l-.001-.122zM17.25 19.128l-.001.144a2.25 2.25 0 01-.233.96 10.088 10.088 0 005.06-1.01.75.75 0 00.42-.643 4.875 4.875 0 00-6.957-4.611 8.586 8.586 0 011.71 5.157v.003z" />
                </svg>
                Job Application Tracker
            </h1>
            """)
        
        # State variables
        base_resume_path = gr.State(None)
        current_job_id = gr.State(None)
        
        with gr.Tabs() as tabs:
            # Tab 1: User Information
            with gr.TabItem("User Information", id="user-info-tab"):
                gr.Markdown('<div class="tab-header">Personal Information</div>')
                
                gr.Markdown('<div class="info-box">This information will be used in your cover letters and other application documents.</div>')
                with gr.Row():
                    with gr.Column():
                        user_name = gr.Textbox(label="Full Name")
                        user_address = gr.Textbox(label="Address", lines=3)
                    
                    with gr.Column():
                        user_phone = gr.Textbox(label="Phone Number")
                        user_email = gr.Textbox(label="Email Address")
                
                with gr.Row():
                    save_user_btn = gr.Button("Save User Information", variant="primary")
                    load_user_btn = gr.Button("Load Saved Information", variant="secondary")
                
                user_info_result = gr.Textbox(label="Result", interactive=False, elem_classes=["result-box"])
                
                # Save user information
                def save_user_info(name, address, phone, email):
                    if not name or not email:
                        return "Please provide at least your name and email address"
                    
                    try:
                        save_user_information(name, address, phone, email)
                        return "User information saved successfully!"
                    except Exception as e:
                        return f"Error saving user information: {str(e)}"
                
                # Load user information
                def load_user_info():
                    user_info = get_user_information()
                    return (
                        user_info.get("full_name", ""),
                        user_info.get("address", ""),
                        user_info.get("phone", ""),
                        user_info.get("email", ""),
                        "User information loaded successfully!"
                    )
                
                # Connect event handlers
                save_user_btn.click(
                    save_user_info,
                    inputs=[user_name, user_address, user_phone, user_email],
                    outputs=[user_info_result]
                )
                
                load_user_btn.click(
                    load_user_info,
                    inputs=[],
                    outputs=[user_name, user_address, user_phone, user_email, user_info_result]
                )
            
            # Tab 2: Add New Application
            with gr.TabItem("Add Application", id="add-app-tab"):
                gr.Markdown('<div class="tab-header">Add New Job Application</div>')
                
                gr.Markdown('<div class="info-box">Track a new job application by filling out the details below.</div>')
                with gr.Row():
                    with gr.Column():
                        company_input = gr.Textbox(label="Company Name")
                        position_input = gr.Textbox(label="Position")
                        date_input = gr.Textbox(label="Date Applied", value=datetime.datetime.now().strftime("%Y-%m-%d"))
                        status_input = gr.Dropdown(
                            label="Status", 
                            choices=["Applied", "Interviewing", "Rejected", "Offer", "Accepted", "Declined"],
                            value="Applied"
                        )
                        salary_input = gr.Textbox(label="Salary Information (Optional)")
                    
                    with gr.Column():
                        contact_input = gr.Textbox(label="Contact Information (Optional)")
                        url_input = gr.Textbox(label="Application URL (Optional)")
                        job_desc_input = gr.Textbox(label="Job Description", lines=10)
                        notes_input = gr.Textbox(label="Notes (Optional)", lines=3)
                
                with gr.Row():
                    base_resume_upload = gr.File(label="Upload Base Resume (DOCX)")
                    
                with gr.Row():
                    add_btn = gr.Button("Add Job Application", variant="primary")
                    clear_btn = gr.Button("Clear Form", variant="secondary")
                
                result_text = gr.Textbox(label="Result", interactive=False, elem_classes=["result-box"])
                
                # Add job application function
                def add_application(company, position, date, job_desc, status, salary, contact, url, notes, resume_file, resume_path_state):
                    print("Starting add_application function...")
                    if not company or not position or not date:
                        print("Missing required fields")
                        return "Please fill in all required fields (Company, Position, Date)", resume_path_state
                    
                    try:
                        # Save the base resume if provided
                        resume_path = None
                        if resume_file:
                            print("Saving uploaded resume...")
                            resume_path = save_uploaded_resume(resume_file)
                            resume_path_state = resume_path
                            print(f"Resume saved to {resume_path}")
                        
                        # Add to database
                        print("Adding job application to database...")
                        job_id = add_job_application(
                            company, position, date, job_desc, status,
                            salary, contact, url, notes, resume_path, None
                        )
                        print(f"Job application added with ID: {job_id}")
                        
                        return f"Job application added successfully! ID: {job_id}", resume_path_state
                    except Exception as e:
                        print(f"Error in add_application: {str(e)}")
                        return f"Error adding job application: {str(e)}", resume_path_state
                
                # Clear form function
                def clear_form():
                    return "", "", datetime.datetime.now().strftime("%Y-%m-%d"), "", "Applied", "", "", "", "", None
                
                # Connect event handlers
                add_btn.click(
                    add_application,
                    inputs=[
                        company_input, position_input, date_input, job_desc_input, status_input,
                        salary_input, contact_input, url_input, notes_input, base_resume_upload, base_resume_path
                    ],
                    outputs=[result_text, base_resume_path]
                )
                
                clear_btn.click(
                    clear_form,
                    inputs=[],
                    outputs=[
                        company_input, position_input, date_input, job_desc_input, status_input,
                        salary_input, contact_input, url_input, notes_input, base_resume_upload
                    ]
                )
            
            # Tab 3: View/Edit Applications
            with gr.TabItem("View Applications", id="view-app-tab"):
                gr.Markdown('<div class="tab-header">Manage Job Applications</div>')
                
                gr.Markdown('<div class="info-box">View, edit, or delete your existing job applications.</div>')
                with gr.Row():
                    refresh_btn = gr.Button("Refresh List", variant="primary", size="sm")
                
                applications_table = gr.DataFrame(
                    headers=['ID', 'Company', 'Position', 'Date Applied', 'Status'],
                    datatype=['number', 'str', 'str', 'str', 'str'],
                    label="Job Applications"
                )
                
                with gr.Row():
                    selected_id_input = gr.Number(label="Selected Job ID", precision=0)
                    load_btn = gr.Button("Load Application", variant="primary", size="sm")
                    delete_btn = gr.Button("Delete Application", variant="stop", size="sm")
                
                with gr.Row():
                    with gr.Column():
                        edit_company = gr.Textbox(label="Company Name")
                        edit_position = gr.Textbox(label="Position")
                        edit_date = gr.Textbox(label="Date Applied")
                        edit_status = gr.Dropdown(
                            label="Status", 
                            choices=["Applied", "Interviewing", "Rejected", "Offer", "Accepted", "Declined"]
                        )
                        edit_salary = gr.Textbox(label="Salary Information")
                    
                    with gr.Column():
                        edit_contact = gr.Textbox(label="Contact Information")
                        edit_url = gr.Textbox(label="Application URL")
                        edit_job_desc = gr.Textbox(label="Job Description", lines=10)
                        edit_notes = gr.Textbox(label="Notes", lines=3)
                
                with gr.Row():
                    update_btn = gr.Button("Update Application", variant="primary")
                
                edit_result = gr.Textbox(label="Result", interactive=False, elem_classes=["result-box"])
                
                # Load applications on refresh
                def load_applications():
                    return refresh_applications_list()
                
                # Load specific application details
                def load_application_details(job_id):
                    if not job_id:
                        return "Please enter a valid Job ID", "", "", "", "", "", "", "", "", None
                    
                    application = get_job_application(int(job_id))
                    if not application:
                        return "Job application not found", "", "", "", "", "", "", "", "", None
                    
                    return (
                        f"Loaded application {job_id}",
                        application["company_name"],
                        application["position"],
                        application["date_applied"],
                        application["status"],
                        application["salary_info"] or "",
                        application["contact_info"] or "",
                        application["application_url"] or "",
                        application["job_description"] or "",
                        application["notes"] or "",
                        job_id
                    )
                
                # Update application
                def update_application(company, position, date, status, salary, contact, url, job_desc, notes, job_id):
                    if not job_id:
                        return "No application selected"
                    
                    update_job_application(
                        job_id,
                        company_name=company,
                        position=position,
                        date_applied=date,
                        status=status,
                        salary_info=salary,
                        contact_info=contact,
                        application_url=url,
                        job_description=job_desc,
                        notes=notes
                    )
                    
                    return f"Application {job_id} updated successfully"
                
                # Delete application
                def delete_application(job_id):
                    if not job_id:
                        return "Please enter a valid Job ID", None
                    
                    success = delete_job_application(int(job_id))
                    if success:
                        return f"Application {job_id} deleted successfully", None
                    else:
                        return f"Failed to delete application {job_id}", job_id
                
                # Connect event handlers
                refresh_btn.click(
                    load_applications,
                    inputs=[],
                    outputs=[applications_table]
                )
                
                load_btn.click(
                    load_application_details,
                    inputs=[selected_id_input],
                    outputs=[
                        edit_result, edit_company, edit_position, edit_date, edit_status,
                        edit_salary, edit_contact, edit_url, edit_job_desc, edit_notes, current_job_id
                    ]
                )
                
                update_btn.click(
                    update_application,
                    inputs=[
                        edit_company, edit_position, edit_date, edit_status, edit_salary,
                        edit_contact, edit_url, edit_job_desc, edit_notes, current_job_id
                    ],
                    outputs=[edit_result]
                )
                
                delete_btn.click(
                    delete_application,
                    inputs=[selected_id_input],
                    outputs=[edit_result, current_job_id]
                )
            
            # Tab 4: Resume Customization
            with gr.TabItem("Customize Resume", id="resume-tab"):
                gr.Markdown('<div class="tab-header">Resume Customization</div>')
                
                gr.Markdown('<div class="info-box">Tailor your resume to match the job description using AI assistance.</div>')
                with gr.Row():
                    with gr.Column():
                        resume_job_id = gr.Number(label="Job ID (Optional)", precision=0)
                        load_job_btn = gr.Button("Load Job Details", variant="primary", size="sm")
                    
                    with gr.Column():
                        resume_upload = gr.File(label="Upload Resume (DOCX)")
                
                resume_job_company = gr.Textbox(label="Company Name", interactive=False)
                resume_job_position = gr.Textbox(label="Position", interactive=False)
                resume_job_desc = gr.Textbox(label="Job Description", lines=10)
                
                customize_btn = gr.Button("Customize Resume", variant="primary")
                
                with gr.Row():
                    resume_result = gr.Textbox(label="Customization Result", interactive=False)
                    resume_download = gr.File(label="Download Customized Resume")
                
                # Load job details for resume customization
                def load_job_for_resume(job_id, resume_path_state):
                    if not job_id:
                        return "Please enter a valid Job ID", "", "", resume_path_state
                    
                    application = get_job_application(int(job_id))
                    if not application:
                        return "Job application not found", "", "", resume_path_state
                    
                    return (
                        f"Loaded job details for {application['company_name']}",
                        application["company_name"],
                        application["position"],
                        application["job_description"] or "",
                        resume_path_state
                    )
                
                # Customize resume
                def customize_resume(job_desc, resume_file, resume_path_state):
                    if not job_desc:
                        return "Please provide a job description", None, resume_path_state
                    
                    # Determine which resume to use
                    resume_path = None
                    if resume_file:
                        resume_path = save_uploaded_resume(resume_file)
                    elif resume_path_state:
                        resume_path = resume_path_state
                    else:
                        return "Please upload a resume", None, resume_path_state
                    
                    # Create customized resume
                    output_path, suggestions = create_custom_resume(resume_path, job_desc)
                    
                    return f"Resume customized successfully!\n\nSuggestions:\n{suggestions}", output_path, resume_path_state
                
                # Connect event handlers
                load_job_btn.click(
                    load_job_for_resume,
                    inputs=[resume_job_id, base_resume_path],
                    outputs=[resume_result, resume_job_company, resume_job_position, resume_job_desc, base_resume_path]
                )
                
                customize_btn.click(
                    customize_resume,
                    inputs=[resume_job_desc, resume_upload, base_resume_path],
                    outputs=[resume_result, resume_download, base_resume_path]
                )
            
            # Tab 5: Cover Letter Generation
            with gr.TabItem("Generate Cover Letter", id="cover-letter-tab"):
                gr.Markdown('<div class="tab-header">Cover Letter Generation</div>')
                
                gr.Markdown('<div class="info-box">Generate a professional cover letter tailored to the job description using AI.</div>')
                with gr.Row():
                    with gr.Column():
                        cl_job_id = gr.Number(label="Job ID (Optional)", precision=0)
                        cl_load_job_btn = gr.Button("Load Job Details", variant="primary", size="sm")
                    
                    with gr.Column():
                        cl_resume_upload = gr.File(label="Upload Resume (DOCX) *Required*")
                
                cl_company = gr.Textbox(label="Company Name")
                cl_position = gr.Textbox(label="Position")
                cl_job_desc = gr.Textbox(label="Job Description", lines=10)
                cl_resume_path = gr.State(None)  # Hidden state to store resume path
                
                generate_cl_btn = gr.Button("Generate Cover Letter", variant="primary")
                
                with gr.Row():
                    cl_result = gr.Textbox(label="Generated Cover Letter Preview", lines=10, interactive=False)
                    cl_download = gr.File(label="Download Cover Letter")
                
                # Load job details for cover letter
                def load_job_for_cl(job_id):
                    if not job_id:
                        return "Please enter a valid Job ID", "", "", "", None
                    
                    application = get_job_application(int(job_id))
                    if not application:
                        return "Job application not found", "", "", "", None
                    
                    # Get resume path if available
                    resume_path = application.get("resume_path")
                    resume_message = ""
                    if resume_path and os.path.exists(resume_path):
                        resume_message = f" (Resume found: {os.path.basename(resume_path)})"
                    else:
                        resume_path = None
                        resume_message = " (No resume found, please upload one)"
                    
                    return (
                        f"Loaded job details for {application['company_name']}{resume_message}",
                        application["company_name"],
                        application["position"],
                        application["job_description"] or "",
                        resume_path
                    )
                
                # Generate cover letter
                def generate_cover_letter_handler(company, position, job_desc, resume_file, resume_path_state):
                    if not company or not position or not job_desc:
                        return "Please fill in all fields", None, resume_path_state
                    
                    # Determine which resume to use
                    resume_path = None
                    if resume_file:
                        resume_path = save_uploaded_resume(resume_file)
                        resume_path_state = resume_path
                    elif resume_path_state:
                        resume_path = resume_path_state
                    else:
                        return "Please upload a resume or load a job with an existing resume", None, resume_path_state
                    
                    output_path, cover_letter_text = generate_cover_letter(job_desc, company, position, resume_path)
                    
                    return cover_letter_text, output_path, resume_path_state
                
                # Connect event handlers
                cl_load_job_btn.click(
                    load_job_for_cl,
                    inputs=[cl_job_id],
                    outputs=[cl_result, cl_company, cl_position, cl_job_desc, cl_resume_path]
                )
                
                generate_cl_btn.click(
                    generate_cover_letter_handler,
                    inputs=[cl_company, cl_position, cl_job_desc, cl_resume_upload, cl_resume_path],
                    outputs=[cl_result, cl_download, cl_resume_path]
                )
        
        # Footer
        gr.Markdown('<div class="footer">Job Application Tracker - Helping you organize your job search</div>')
        
        # Initialize the database on startup
        def init_app():
            setup_database()
            return None
        
        app.load(init_app, inputs=None, outputs=None)
    
    return app

# Main function
def main():
    app = create_ui()
    app.launch(
        share=False,
    )

if __name__ == "__main__":
    main()
