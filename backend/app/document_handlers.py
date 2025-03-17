import os
from docx import Document
from docx.shared import Pt
import datetime
from fastapi import HTTPException

from app.ai_integration import generate_resume_suggestions, generate_cover_letter
from app.file_handler import get_document_text
from app.database import get_user_information

# Constants
RESUME_FOLDER = "resumes"
COVER_LETTER_FOLDER = "cover_letters"

def create_custom_resume(base_resume_path, job_description, company_name=None, position=None, output_path=None):
    """
    Create a customized resume based on job description using AI suggestions
    
    Args:
        base_resume_path: Path to the base resume file
        job_description: Job description to customize for
        company_name: Company name for file naming (optional)
        position: Position title for file naming (optional)
        output_path: Optional custom output path
        
    Returns:
        tuple: (output_path, tailoring_suggestions)
    """
    # Debug output
    print(f"DEBUG - create_custom_resume called with:")
    print(f"  base_resume_path: {base_resume_path}")
    print(f"  job_description length: {len(job_description) if job_description else 0}")
    print(f"  output_path: {output_path}")
    print(f"  File exists: {os.path.exists(base_resume_path)}")
    
    # If using relative path, make sure it's relative to current working directory
    if not os.path.isabs(base_resume_path) and not base_resume_path.startswith(RESUME_FOLDER):
        adjusted_path = os.path.join(RESUME_FOLDER, os.path.basename(base_resume_path))
        print(f"  Adjusted path: {adjusted_path}")
        if os.path.exists(adjusted_path):
            base_resume_path = adjusted_path
            print(f"  Using adjusted path: {base_resume_path}")
    
    if not os.path.exists(base_resume_path):
        raise HTTPException(status_code=404, detail=f"Resume file not found at path: {base_resume_path}")
    
    # Generate tailoring suggestions using Ollama
    tailoring_suggestions = generate_resume_suggestions(job_description)
    
    try:
        # Load the base resume
        doc = Document(base_resume_path)
        
        # Add tailoring suggestions at the end
        # Check if the style exists in the document
        try:
            doc.add_heading("Tailoring Suggestions", level=1)
        except KeyError:
            # If the style doesn't exist, use a paragraph with manual formatting
            p = doc.add_paragraph("Tailoring Suggestions")
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)  # Approximate size for Heading 1
        
        doc.add_paragraph(tailoring_suggestions)
        
        # Save the customized resume
        if not output_path:
            # Create a formatted filename
            if company_name and position:
                # Format today's date as YYYY-MM-DD
                today_formatted = datetime.datetime.now().strftime("%Y-%m-%d")
                
                # Clean up file name to ensure it's valid (replace invalid characters)
                clean_company = company_name.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')
                clean_position = position.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')
                
                # Create filename with em dash separators
                filename = f"{clean_company}—{clean_position}—Resume—{today_formatted}.docx"
            else:
                # Use timestamp based naming if company/position not provided
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"custom_resume_{timestamp}.docx"
                
            output_path = os.path.join(RESUME_FOLDER, filename)
            
            # Ensure directory exists
            os.makedirs(RESUME_FOLDER, exist_ok=True)
        
        doc.save(output_path)
        return output_path, tailoring_suggestions
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating custom resume: {str(e)}")

def create_cover_letter(job_description, company_name, position, resume_path, output_path=None):
    """
    Create a cover letter based on job description, company, position and resume
    
    Args:
        job_description: Job description text
        company_name: Company name
        position: Position title
        resume_path: Path to the resume file
        output_path: Optional custom output path
        
    Returns:
        tuple: (output_path, cover_letter_text)
    """
    # Get user information
    user_info = get_user_information()
    print(f"DEBUG - User info retrieved: {user_info}")
    # Debug output
    print(f"DEBUG - create_cover_letter called with:")
    print(f"  resume_path: {resume_path}")
    print(f"  company_name: {company_name}")
    print(f"  position: {position}")
    print(f"  job_description length: {len(job_description) if job_description else 0}")
    print(f"  output_path: {output_path}")
    print(f"  File exists: {os.path.exists(resume_path)}")
    
    # If using relative path, make sure it's relative to current working directory
    if not os.path.isabs(resume_path) and not resume_path.startswith(RESUME_FOLDER):
        adjusted_path = os.path.join(RESUME_FOLDER, os.path.basename(resume_path))
        print(f"  Adjusted path: {adjusted_path}")
        if os.path.exists(adjusted_path):
            resume_path = adjusted_path
            print(f"  Using adjusted path: {resume_path}")
    
    if not os.path.exists(resume_path):
        raise HTTPException(status_code=404, detail=f"Resume file not found at path: {resume_path}")
    
    # Extract resume content
    resume_text = get_document_text(resume_path)
    
    # Generate cover letter using Ollama
    # The function will automatically pull user profile information as needed
    cover_letter_text = generate_cover_letter(
        job_description, 
        company_name, 
        position, 
        resume_text
    )
    
    try:
        # Create a new document
        doc = Document()
        
        # Add the user's contact information at the top
        if user_info.get('full_name'):
            doc.add_paragraph(user_info.get('full_name'))
        if user_info.get('address'):
            doc.add_paragraph(user_info.get('address'))
        if user_info.get('phone'):
            doc.add_paragraph(user_info.get('phone'))
        if user_info.get('email'):
            doc.add_paragraph(user_info.get('email'))
        
        doc.add_paragraph()  # Add space after contact info
        
        # Add today's date in the format "Month Day, Year"
        today_date = datetime.datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(today_date)
        doc.add_paragraph()
        
        # Add company info placeholders
        doc.add_paragraph("Hiring Manager")
        doc.add_paragraph(f"{company_name}")
        doc.add_paragraph("Company Address")
        doc.add_paragraph("City, State ZIP")
        doc.add_paragraph()
        
        # Add greeting
        doc.add_paragraph("Dear Hiring Manager,")
        
        # Add cover letter content
        paragraphs = cover_letter_text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Add closing with user's name if available
        doc.add_paragraph()
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        if user_info.get('full_name'):
            doc.add_paragraph(user_info.get('full_name'))
        else:
            doc.add_paragraph("[Your Name]")
        
        # Save the cover letter with the format: CompanyName—Position—Date.docx
        if not output_path:
            # Format today's date as YYYY-MM-DD
            today_formatted = datetime.datetime.now().strftime("%Y-%m-%d")
            
            # Clean up file name to ensure it's valid (replace invalid characters)
            clean_company = company_name.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')
            clean_position = position.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')
            
            # Create filename with em dash separators
            filename = f"{clean_company}—{clean_position}—{today_formatted}.docx"
            output_path = os.path.join(COVER_LETTER_FOLDER, filename)
            
            # Ensure directory exists
            os.makedirs(COVER_LETTER_FOLDER, exist_ok=True)
        
        doc.save(output_path)
        return output_path, cover_letter_text
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating cover letter: {str(e)}")
